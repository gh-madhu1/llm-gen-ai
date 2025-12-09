"""
Document generator module for creating formatted DOCX files.
Handles all document creation and formatting logic.
"""
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from llm_gen_ai.utils import sanitize_filename, clean_content
from llm_gen_ai.config import SECTION_ORDER, OUTPUT_DIR, MAX_FILENAME_LENGTH


class DocumentGenerator:
    """Handles DOCX document creation with proper formatting."""

    def __init__(self, output_dir=OUTPUT_DIR):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def create_document(self, paper_content, filename="white_paper.docx"):
        """
        Create a DOCX document from paper content.
        
        Args:
            paper_content: Dictionary of section_name -> content
            filename: Output filename
        
        Returns:
            Path to created document
        """
        # Sanitize filename
        filename = sanitize_filename(filename, max_length=MAX_FILENAME_LENGTH)
        full_path = self._get_versioned_path(filename)

        # Create document
        doc = Document()
        
        # Add title section
        self._add_title_section(doc, paper_content)
        
        # Add author placeholder
        self._add_author_section(doc)
        
        # Add content sections in order
        self._add_content_sections(doc, paper_content)
        
        # Save document
        doc.save(full_path)
        print(f"\n{'='*60}")
        print(f"✅ White paper saved to: {full_path}")
        print(f"{'='*60}\n")
        
        return full_path

    def _get_versioned_path(self, filename):
        """
        Get versioned file path to avoid overwriting.
        
        Args:
            filename: Desired filename
        
        Returns:
            Versioned file path
        """
        base_name = os.path.splitext(filename)[0]
        extension = os.path.splitext(filename)[1]
        full_path = os.path.join(self.output_dir, filename)

        # Check if file exists and create versioned filename
        version = 1
        while os.path.exists(full_path):
            version += 1
            full_path = os.path.join(
                self.output_dir, f"{base_name}_v{version}{extension}"
            )
        
        return full_path

    def _add_title_section(self, doc, paper_content):
        """Add title to document."""
        if 'title' in paper_content:
            title = doc.add_heading(
                clean_content(paper_content['title']), 0
            )
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()  # Spacing

    def _add_author_section(self, doc):
        """Add author placeholder section."""
        author_para = doc.add_paragraph()
        author_para.add_run('Author: [Author Name]').bold = True
        author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        affiliation_para = doc.add_paragraph()
        affiliation_para.add_run('Affiliation: [Organization/Institution]').italic = True
        affiliation_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        contact_para = doc.add_paragraph()
        contact_para.add_run('Contact: [email@example.com]').italic = True
        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph()  # Spacing
        doc.add_paragraph()  # Spacing

    def _add_content_sections(self, doc, paper_content):
        """
        Add all content sections in proper order.
        
        Args:
            doc: Document object
            paper_content: Dictionary of section content
        """
        added_sections = set(['title'])  # Already added title

        # Add sections in standard order
        for ordered_section in SECTION_ORDER:
            for section_title, content in paper_content.items():
                section_lower = section_title.lower()
                
                if section_lower == ordered_section and section_lower not in added_sections:
                    self._add_section(doc, section_title, content)
                    added_sections.add(section_lower)

        # Add any remaining sections not in standard order
        for section_title, content in paper_content.items():
            if section_title.lower() not in added_sections:
                self._add_section(doc, section_title, content)

    def _add_section(self, doc, section_title, content):
        """
        Add a single section to the document.
        
        Args:
            doc: Document object
            section_title: Title of the section
            content: Section content
        """
        # Clean content before adding
        clean_section_content = clean_content(content)
        
        if clean_section_content:  # Only add if there's content
            doc.add_heading(section_title, level=1)
            
            # Handle multi-paragraph content
            paragraphs = clean_section_content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            doc.add_paragraph()  # Spacing between sections
