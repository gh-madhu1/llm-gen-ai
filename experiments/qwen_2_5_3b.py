import time
import re
import os
from functools import wraps
from transformers import AutoTokenizer, AutoModelForCausalLM
import torch
from duckduckgo_search import DDGS
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

if torch.backends.mps.is_available():
    device = torch.device("mps")
else:
    device = torch.device("cpu")

# Tracks the process execution time
def track_process_time(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        start_time = time.process_time()
        result = func(*args, **kwargs)
        end_time = time.process_time()
        process_time = end_time - start_time
        print(f"Process time for {func.__name__}: {process_time:.3f} seconds")
        return result
    return wrapper

@track_process_time
def load_pretrained_model(model_path):
    # Load the pre-trained model onto the selected device.
    model = AutoModelForCausalLM.from_pretrained(model_path, device_map=device)
    return model

@track_process_time
def load_tokenizer_model(model_path):
    return AutoTokenizer.from_pretrained(model_path)

class WhitePaperAgent:
    def __init__(self, model, tokenizer, output_dir="."):
        self.model = model
        self.tokenizer = tokenizer
        self.ddgs = DDGS()
        self.max_steps = 20  # Increased for complete white paper generation
        self.history = []
        self.output_dir = output_dir
        self.paper_content = {}

    def search(self, query):
        print(f"Searching for: {query}")
        results = self.ddgs.text(query, max_results=3)
        if results:
            return "\n".join([f"- {r['title']}: {r['body']}" for r in results])
        return "No results found."
    
    def sanitize_filename(self, filename):
        """Sanitize filename to prevent errors and limit length."""
        # Remove any special characters and limit length
        filename = filename.strip()
        # Remove path separators and other problematic characters
        filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        # Limit filename length (excluding extension)
        base_name = os.path.splitext(filename)[0]
        extension = os.path.splitext(filename)[1] if '.' in filename else '.docx'
        if len(base_name) > 50:
            base_name = base_name[:50]
        return f"{base_name}{extension}"
    
    def clean_content(self, text):
        """Remove agent thoughts and actions from content."""
        # Remove THOUGHT, ACTION, OBSERVATION patterns
        text = re.sub(r'THOUGHT:.*?(?=\n|$)', '', text, flags=re.IGNORECASE)
        text = re.sub(r'ACTION:.*?(?=\n|$)', '', text, flags=re.IGNORECASE)
        text = re.sub(r'OBSERVATION:.*?(?=\n|$)', '', text, flags=re.IGNORECASE)
        # Clean up extra whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        return text.strip()
    
    def write_to_docx(self, filename="white_paper.docx"):
        """Write the collected paper content to a DOCX file with versioning."""
        # Sanitize filename
        filename = self.sanitize_filename(filename)
        base_name = os.path.splitext(filename)[0]
        extension = os.path.splitext(filename)[1]
        full_path = os.path.join(self.output_dir, filename)
        
        # Check if file exists and create versioned filename
        version = 1
        while os.path.exists(full_path):
            version += 1
            full_path = os.path.join(self.output_dir, f"{base_name}_v{version}{extension}")
        
        # Create document
        doc = Document()
        
        # Define standard white paper section order
        section_order = [
            'title', 'author', 'abstract', 'executive summary', 'introduction',
            'background', 'problem statement', 'methodology', 'solution',
            'implementation', 'results', 'benefits', 'challenges',
            'recommendations', 'conclusion', 'references', 'appendix'
        ]
        
        # Add title
        if 'title' in self.paper_content:
            title = doc.add_heading(self.clean_content(self.paper_content['title']), 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()  # Add spacing
        
        # Add author placeholder
        author_para = doc.add_paragraph()
        author_para.add_run('Author: [Author Name]').bold = True
        author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        affiliation_para = doc.add_paragraph()
        affiliation_para.add_run('Affiliation: [Organization/Institution]').italic = True
        affiliation_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        contact_para = doc.add_paragraph()
        contact_para.add_run('Contact: [email@example.com]').italic = True
        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()  # Add spacing
        doc.add_paragraph()  # Add spacing
        
        # Add sections in proper order
        added_sections = set(['title'])  # Already added title
        for ordered_section in section_order:
            for section_title, content in self.paper_content.items():
                if section_title.lower() == ordered_section and section_title.lower() not in added_sections:
                    # Clean content before adding
                    clean_section_content = self.clean_content(content)
                    if clean_section_content:  # Only add if there's content after cleaning
                        doc.add_heading(section_title, level=1)
                        # Handle multi-paragraph content
                        paragraphs = clean_section_content.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                doc.add_paragraph(para.strip())
                        doc.add_paragraph()  # Add spacing between sections
                    added_sections.add(section_title.lower())
        
        # Add any remaining sections not in the standard order
        for section_title, content in self.paper_content.items():
            if section_title.lower() not in added_sections:
                clean_section_content = self.clean_content(content)
                if clean_section_content:
                    doc.add_heading(section_title, level=1)
                    paragraphs = clean_section_content.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
                    doc.add_paragraph()
        
        # Save document
        doc.save(full_path)
        print(f"\n{'='*60}")
        print(f"✅ White paper saved to: {full_path}")
        print(f"{'='*60}\n")
        return full_path

    def generate_text(self, prompt, max_new_tokens=1024):
        inputs = self.tokenizer(prompt, return_tensors="pt").to(device)
        generated = self.model.generate(
            inputs.input_ids, 
            max_new_tokens=max_new_tokens,
            do_sample=True,
            temperature=0.7
        )
        return self.tokenizer.decode(generated[0], skip_special_tokens=True)[len(prompt):]

    def reason_and_act(self, idea):
        print(f"Starting research on: {idea}")
        self.history = []
        
        system_prompt = f"""You are an AI agent designed to write a COMPREHENSIVE, PUBLICATION-READY white paper based on:

"{idea}"

Your goal is to create a complete white paper with ALL of the following sections:

1. Title (use WRITE[title | Your Title Here])
2. Abstract (150-250 words summary)
3. Executive Summary (key findings and recommendations)
4. Introduction (context, importance, scope)
5. Background (relevant history and context)
6. Problem Statement (clear definition of the issue)
7. Solution/Methodology (your proposed approach)
8. Implementation (how to implement the solution)
9. Benefits (advantages and value proposition)
10. Challenges and Risks (potential issues and mitigation)
11. Recommendations (actionable next steps)
12. Conclusion (summary and future outlook)

You have access to the following tools:

1. SEARCH: Search the web for information. Usage: SEARCH[query]
2. WRITE: Write a section of the white paper. Usage: WRITE[section_title | content]
3. SAVE_DOCX: Save the white paper to a DOCX file. Usage: SAVE_DOCX[filename]
4. FINISH: Complete the task. Usage: FINISH[summary]

IMPORTANT GUIDELINES:
- Research thoroughly using SEARCH before writing each major section
- Write detailed, well-researched content (minimum 200-300 words per section)
- Use professional, academic tone
- Include specific examples, data, and evidence
- Ensure logical flow between sections
- Only call FINISH after ALL sections are written

Format your response as:
THOUGHT: [Your reasoning]
ACTION: [One of the tools]

Begin by researching the topic!
"""
        current_prompt = system_prompt
        
        for step in range(self.max_steps):
            print(f"\n--- Step {step + 1} ---")
            response = self.generate_text(current_prompt, max_new_tokens=200)
            print(f"Agent Response:\n{response}")
            
            # Parse response
            thought_match = re.search(r"THOUGHT:(.*?)(?=ACTION:|$)", response, re.DOTALL)
            action_match = re.search(r"ACTION:(.*)", response, re.DOTALL)
            
            if thought_match:
                thought = thought_match.group(1).strip()
                self.history.append(f"THOUGHT: {thought}")
            
            if action_match:
                action_line = action_match.group(1).strip()
                self.history.append(f"ACTION: {action_line}")
                
                if "SEARCH[" in action_line:
                    # Extract search query
                    match = re.search(r'SEARCH\[(.+?)\]', action_line)
                    if match:
                        query = match.group(1)
                        observation = self.search(query)
                        print(f"Observation: {observation[:100]}...")
                        current_prompt += f"\nTHOUGHT: {thought}\nACTION: {action_line}\nOBSERVATION: {observation}\n"
                    else:
                        print("Invalid SEARCH format")
                        current_prompt += f"\nOBSERVATION: Invalid SEARCH format. Use SEARCH[query]\n"
                
                elif "WRITE[" in action_line:
                    # Extract content using regex to handle multi-line
                    match = re.search(r'WRITE\[(.+?)\]', action_line, re.DOTALL)
                    if match:
                        content = match.group(1)
                        if "|" in content:
                            parts = content.split("|", 1)
                            section_title = parts[0].strip()
                            section_content = parts[1].strip() if len(parts) > 1 else ""
                            # Store only the clean content, no agent thoughts
                            self.paper_content[section_title] = section_content
                            print(f"✍️  Writing section: {section_title}")
                            current_prompt += f"\nTHOUGHT: {thought}\nACTION: {action_line}\nOBSERVATION: Section '{section_title}' written successfully.\n"
                        else:
                            print(f"⚠️  Invalid WRITE format (missing |): {content[:50]}...")
                            current_prompt += f"\nOBSERVATION: Invalid WRITE format. Use WRITE[section_title | content]\n"
                    else:
                        print("⚠️  Could not parse WRITE action")
                        current_prompt += f"\nOBSERVATION: Could not parse WRITE action.\n"
                
                elif "SAVE_DOCX[" in action_line:
                    # Extract filename using regex
                    match = re.search(r'SAVE_DOCX\[(.+?)\]', action_line)
                    if match:
                        filename = match.group(1).strip()
                        try:
                            file_path = self.write_to_docx(filename)
                            observation = f"White paper saved to {file_path}"
                            print(observation)
                            current_prompt += f"\nTHOUGHT: {thought}\nACTION: {action_line}\nOBSERVATION: {observation}\n"
                        except Exception as e:
                            observation = f"Error saving file: {str(e)}"
                            print(observation)
                            current_prompt += f"\nOBSERVATION: {observation}\n"
                    else:
                        print("Invalid SAVE_DOCX format")
                        current_prompt += f"\nOBSERVATION: Invalid SAVE_DOCX format. Use SAVE_DOCX[filename]\n"
                
                elif "FINISH[" in action_line:
                    print("Agent finished task.")
                    # Check if we have enough content
                    minimum_sections = 5
                    if len(self.paper_content) < minimum_sections:
                        observation = f"Warning: Only {len(self.paper_content)} sections written. A complete white paper needs at least {minimum_sections} sections. Consider adding more content."
                        print(observation)
                        current_prompt += f"\nTHOUGHT: {thought}\nACTION: {action_line}\nOBSERVATION: {observation}\n"
                    else:
                        # Save to docx before finishing
                        if self.paper_content:
                            self.write_to_docx()
                        print(f"\n✅ White paper complete with {len(self.paper_content)} sections!")
                        return "White paper generation complete."
                
                else:
                    print("Unknown action.")
                    current_prompt += f"\nTHOUGHT: {thought}\nACTION: {action_line}\nOBSERVATION: Invalid action format. Use SEARCH, WRITE, SAVE_DOCX, or FINISH.\n"
            else:
                print("No action found.")
                current_prompt += "\nOBSERVATION: Please specify an ACTION.\n"

        return "Max steps reached."

@track_process_time
def generate(prompt):
    # This function is now a wrapper to start the agent
    agent = WhitePaperAgent(model, tokenizer)
    return agent.reason_and_act(prompt)

if __name__ == "__main__":
    model_path = "Qwen/Qwen2.5-3B-Instruct"
    model = load_pretrained_model(model_path)
    tokenizer = load_tokenizer_model(model_path)
    # model.eval() # Optional, depending on usage

    # You can now provide a detailed idea instead of just a topic
    idea = """
    Explore how best agent/model to surve the right request, Focus on:

1) The agent will choose the best model based on the feedback given to each response when we routed the request to the different models the same query. Based on self agent as a judge rate and human feedback the models/agents which performing well should be used to repond to the queries.

2) This helps the best accurate results surved to the user. Helps the trustworthy of the AI Usage in realtime usecases.

3) You can write a paper explaning how we can implement and achieve the solution

4) Highlight any potential risks and issues.

5) Write how it can help optimise the costs, and highlight if any other optimisation thoughts.
    """
    
    print(generate(idea))