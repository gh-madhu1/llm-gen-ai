import time
import re
import os
from functools import wraps
from transformers import AutoTokenizer, AutoModelForCausalLM
import torch
from ddgs import DDGS
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

if torch.backends.mps.is_available():
    device = torch.device("mps")
else:
    device = torch.device("cpu")

# Tracks the process execution time
def track_process_time(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        start_time = time.process_time()
        result = func(*args, **kwargs)
        end_time = time.process_time()
        process_time = end_time - start_time
        print(f"Process time for {func.__name__}: {process_time:.3f} seconds")
        return result
    return wrapper

@track_process_time
def load_pretrained_model(model_path):
    # Load the pre-trained model onto the selected device.
    model = AutoModelForCausalLM.from_pretrained(model_path, device_map=device)
    return model

@track_process_time
def load_tokenizer_model(model_path):
    return AutoTokenizer.from_pretrained(model_path)

class ContextMemory:
    """Manages context to prevent memory overflow and maintain coherence."""
    
    def __init__(self, max_actions=10):
        self.completed_sections = set()
        self.research_summary = {}
        self.recent_actions = []
        self.max_actions = max_actions
        self.plan = None  # Store the agent's plan
        
    def add_section(self, section_name):
        """Track completed section."""
        self.completed_sections.add(section_name.lower())
    
    def add_research(self, query, findings):
        """Store summarized research."""
        # Keep only key points (first 200 chars)
        self.research_summary[query] = findings[:200]
    
    def add_action(self, action_type, details):
        """Add action to sliding window."""
        self.recent_actions.append((action_type, details))
        # Keep only last N actions
        if len(self.recent_actions) > self.max_actions:
            self.recent_actions.pop(0)
    
    def set_plan(self, plan):
        """Store the agent's plan."""
        self.plan = plan
    
    def get_context_summary(self):
        """Generate concise context summary."""
        summary = []
        
        # Show plan if exists
        if self.plan:
            summary.append(f"📋 PLAN: {self.plan[:200]}...")
        
        # Completed sections
        if self.completed_sections:
            summary.append(f"\n✅ Completed sections ({len(self.completed_sections)}): {', '.join(sorted(self.completed_sections))}")
        
        # Required sections for progress tracking
        required = {'title', 'abstract', 'introduction', 'solution', 'benefits', 'challenges', 'conclusion'}
        remaining = required - self.completed_sections
        if remaining:
            summary.append(f"\n⏳ Still needed: {', '.join(sorted(remaining))}")
        
        # Research summary
        if self.research_summary:
            summary.append(f"\n📚 Research topics covered: {', '.join(list(self.research_summary.keys())[:3])}")
        
        # Recent actions
        if self.recent_actions:
            summary.append("\nRecent actions:")
            for action_type, details in self.recent_actions[-5:]:
                summary.append(f"- {action_type}: {details[:50]}")
        
        return "\n".join(summary)
    
    def should_skip_section(self, section_name):
        """Check if section already completed."""
        return section_name.lower() in self.completed_sections

class WhitePaperAgent:
    def __init__(self, model, tokenizer, output_dir="."):
        self.model = model
        self.tokenizer = tokenizer
        self.ddgs = DDGS()
        self.max_steps = 100  # Increased for complete white paper generation
        self.history = []
        self.output_dir = output_dir
        self.paper_content = {}
        # Initialize KV cache for faster generation
        self.past_key_values = None
        # Citation tracking
        self.citations = []
        self.existing_work = []  # Store validation results for comparative analysis
        # Context memory for intelligent memory management
        self.context_memory = ContextMemory(max_actions=10)

    def search(self, query):
        print(f"Searching for: {query}")
        results = self.ddgs.text(query, max_results=3)
        if results:
            # Track citations
            for result in results:
                citation = {
                    'title': result.get('title', 'Unknown'),
                    'url': result.get('href', ''),
                    'snippet': result.get('body', '')[:200]
                }
                # Avoid duplicates
                if citation not in self.citations:
                    self.citations.append(citation)
            
            findings = "\n".join([f"- {r['title']}: {r['body']}" for r in results])
            # Store in context memory
            self.context_memory.add_research(query, findings)
            self.context_memory.add_action("SEARCH", query)
            return findings
        return "No results found."
    
    def sanitize_filename(self, filename):
        """Sanitize filename to prevent errors and limit length."""
        # Remove any special characters and limit length
        filename = filename.strip()
        # Remove path separators and other problematic characters
        filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        # Limit filename length (excluding extension)
        base_name = os.path.splitext(filename)[0]
        extension = os.path.splitext(filename)[1] if '.' in filename else '.docx'
        if len(base_name) > 50:
            base_name = base_name[:50]
        return f"{base_name}{extension}"
    
    def clean_content(self, text):
        """Remove agent thoughts and actions from content."""
        # Remove THOUGHT, ACTION, OBSERVATION patterns
        text = re.sub(r'THOUGHT:.*?(?=\n|$)', '', text, flags=re.IGNORECASE)
        text = re.sub(r'ACTION:.*?(?=\n|$)', '', text, flags=re.IGNORECASE)
        text = re.sub(r'OBSERVATION:.*?(?=\n|$)', '', text, flags=re.IGNORECASE)
        # Clean up extra whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        return text.strip()
    
    def write_to_docx(self, filename="white_paper.docx"):
        """Write the collected paper content to a DOCX file with versioning."""
        # Sanitize filename
        filename = self.sanitize_filename(filename)
        base_name = os.path.splitext(filename)[0]
        extension = os.path.splitext(filename)[1]
        full_path = os.path.join(self.output_dir, filename)
        
        # Check if file exists and create versioned filename
        version = 1
        while os.path.exists(full_path):
            version += 1
            full_path = os.path.join(self.output_dir, f"{base_name}_v{version}{extension}")
        
        # Create document
        doc = Document()
        
        # Define standard white paper section order
        section_order = [
            'title', 'author', 'abstract', 'executive summary', 'introduction',
            'background', 'problem statement', 'methodology', 'solution',
            'implementation', 'results', 'benefits', 'challenges',
            'recommendations', 'conclusion', 'references', 'appendix'
        ]
        
        # Add title
        if 'title' in self.paper_content:
            title = doc.add_heading(self.clean_content(self.paper_content['title']), 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()  # Add spacing
        
        # Add author placeholder
        author_para = doc.add_paragraph()
        author_para.add_run('Author: [Author Name]').bold = True
        author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        affiliation_para = doc.add_paragraph()
        affiliation_para.add_run('Affiliation: [Organization/Institution]').italic = True
        affiliation_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        contact_para = doc.add_paragraph()
        contact_para.add_run('Contact: [email@example.com]').italic = True
        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()  # Add spacing
        doc.add_paragraph()  # Add spacing
        
        # Add sections in proper order
        added_sections = set(['title'])  # Already added title
        for ordered_section in section_order:
            for section_title, content in self.paper_content.items():
                if section_title.lower() == ordered_section and section_title.lower() not in added_sections:
                    # Clean content before adding
                    clean_section_content = self.clean_content(content)
                    if clean_section_content:  # Only add if there's content after cleaning
                        doc.add_heading(section_title, level=1)
                        # Handle multi-paragraph content
                        paragraphs = clean_section_content.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                doc.add_paragraph(para.strip())
                        doc.add_paragraph()  # Add spacing between sections
                    added_sections.add(section_title.lower())
        
        # Add any remaining sections not in the standard order
        for section_title, content in self.paper_content.items():
            if section_title.lower() not in added_sections:
                clean_section_content = self.clean_content(content)
                if clean_section_content:
                    doc.add_heading(section_title, level=1)
                    paragraphs = clean_section_content.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
                    doc.add_paragraph()
        
        # Save document
        doc.save(full_path)
        print(f"\n{'='*60}")
        print(f"✅ White paper saved to: {full_path}")
        print(f"{'='*60}\n")
        return full_path

    def generate_text(self, prompt, max_new_tokens=100, use_cache=True):
        """Generate text with KV caching for faster inference."""
        # Truncate prompt if too long to prevent OOM
        max_prompt_length = 8000  # characters, not tokens
        if len(prompt) > max_prompt_length:
            print(f"\n⚠️  Prompt too long ({len(prompt)} chars), truncating...")
            # Keep the system prompt and recent history
            prompt = prompt[:1000] + "...\n" + prompt[-3000:]
        
        inputs = self.tokenizer(prompt, return_tensors="pt", truncation=True, max_length=2048).to(device)
        
        # Clear MPS cache to prevent OOM
        if device.type == "mps":
            torch.mps.empty_cache()
        
        # Enable KV caching for faster generation
        with torch.no_grad():  # Disable gradients for inference
            generated = self.model.generate(
                inputs.input_ids,
                attention_mask=inputs["attention_mask"],
                max_new_tokens=max_new_tokens,
                do_sample=False,  # Greedy decoding for speed
                use_cache=use_cache,  # Enable KV cache
                pad_token_id=self.tokenizer.pad_token_id or self.tokenizer.eos_token_id
            )
        
        # Clear cache after generation
        if device.type == "mps":
            torch.mps.empty_cache()
        
        return self.tokenizer.decode(generated[0], skip_special_tokens=True)[len(prompt):]

    def validate_idea_novelty(self, idea):
        """Check if the idea already exists on the internet using AI analysis."""
        print("\n" + "="*60)
        print("🔍 STEP 1: Validating Idea Novelty (AI-Powered)")
        print("="*60)
        
        # Create a concise search query from the idea
        idea_summary = idea[:200].replace('\n', ' ')
        search_query = f"{idea_summary} research paper white paper"
        
        print(f"Searching: {search_query[:100]}...")
        
        try:
            results = self.ddgs.text(search_query, max_results=5)
            
            if results and len(results) >= 2:
                # Store all results for potential comparative analysis
                for result in results[:5]:
                    self.existing_work.append({
                        'title': result.get('title', 'Unknown'),
                        'url': result.get('href', ''),
                        'summary': result.get('body', '')[:300]
                    })
                
                print("\n🤖 Analyzing search results with AI...")
                
                # Prepare results summary for AI analysis
                results_text = "\n".join([
                    f"{i+1}. {r['title']}: {r['body'][:200]}" 
                    for i, r in enumerate(results[:5])
                ])
                
                # Use the agent to analyze relevancy and novelty
                analysis_prompt = f"""Analyze if this idea is novel:

PROPOSED IDEA:
{idea[:500]}

EXISTING SEARCH RESULTS:
{results_text}

Determine:
1. Are these search results RELEVANT to the proposed idea? (yes/no)
2. If relevant, does the idea already exist or is it novel? (exists/novel)

Respond ONLY with one of:
- IRRELEVANT: Search results not related to the idea
- EXISTS: Idea clearly already exists in similar form
- NOVEL: Idea appears novel despite similar work

Response:"""
                
                # Get AI analysis
                analysis = self.generate_text(analysis_prompt, max_new_tokens=50).strip().upper()
                
                print(f"📊 AI Analysis Result: {analysis[:100]}")
                
                # Parse AI decision
                if "IRRELEVANT" in analysis:
                    print("\n✅ SEARCH RESULTS IRRELEVANT")
                    print("="*60)
                    print("AI determined search results are not related to your idea.")
                    print("Proceeding with paper generation...")
                    print("="*60)
                    return True
                    
                elif "EXISTS" in analysis:
                    print("\n⚠️  SIMILAR WORK EXISTS")
                    print("="*60)
                    print("AI found similar existing research:")
                    for i, result in enumerate(results[:3], 1):
                        print(f"\n{i}. {result['title']}")
                        print(f"   {result['body'][:150]}...")
                    print("\n" + "="*60)
                    print("❌ CONCLUSION: Similar work already exists.")
                    print("💡 RECOMMENDATION: Refine your idea to be more specific or novel.")
                    print("="*60)
                    return False
                    
                else:  # NOVEL or unclear
                    print("\n✅ IDEA APPEARS NOVEL")
                    print("="*60)
                    print("AI analysis: Idea is sufficiently novel despite existing work.")
                    print("Proceeding with research...")
                    print("="*60)
                    return True
            else:
                # Very few or no results
                if results:
                    for result in results:
                        self.existing_work.append({
                            'title': result.get('title', 'Unknown'),
                            'url': result.get('href', ''),
                            'summary': result.get('body', '')[:300]
                        })
                
                print("\n✅ MINIMAL EXISTING CONTENT")
                print("="*60)
                print("Very few search results found. Idea appears novel.")
                print("="*60)
                return True
                
        except Exception as e:
            print(f"\n⚠️  Validation error: {e}")
            print("Proceeding with caution...\n")
            return True  # Proceed if validation fails
    
    def generate_comparative_analysis(self):
        """Generate a comparative analysis section based on existing work."""
        if not self.existing_work:
            return "No existing work found for comparison."
        
        analysis = "This paper distinguishes itself from existing research in several key ways:\n\n"
        
        for i, work in enumerate(self.existing_work[:3], 1):
            analysis += f"{i}. **{work['title']}**: While this work addresses related topics, "
            analysis += f"our approach differs by focusing on novel aspects and providing unique insights.\n\n"
        
        analysis += "Our contribution provides a fresh perspective by combining theoretical frameworks with practical implementation strategies."
        return analysis
    
    def generate_references(self):
        """Generate a properly formatted references section."""
        if not self.citations:
            return "No references available."
        
        refs = "# References\n\n"
        for i, citation in enumerate(self.citations, 1):
            refs += f"{i}. **{citation['title']}**\n"
            if citation['url']:
                refs += f"   URL: {citation['url']}\n"
            refs += f"   Accessed: {time.strftime('%Y-%m-%d')}\n\n"
        
        return refs
    
    def reason_and_act(self, idea):
        # Step 1: Validate idea novelty
        if not self.validate_idea_novelty(idea):
            return "Idea validation failed: Similar content already exists. Please refine your idea."
        
        # Step 2: Auto-generate comparative analysis if we have existing work
        if self.existing_work:
            comparative_analysis = self.generate_comparative_analysis()
            self.paper_content['Related Work'] = comparative_analysis
            print("\n📊 Generated comparative analysis section")
        
        # Step 3: Proceed with white paper generation
        print(f"\n" + "="*60)
        print("🚀 STEP 2: Starting White Paper Generation")
        print("="*60)
        print(f"Topic: {idea[:100]}...\n")
        self.history = []
        
        system_prompt = f"""You are an AI agent designed to write a COMPREHENSIVE, PUBLICATION-READY white paper based on:

"{idea}"

Your goal is to create a complete white paper with ALL of the following sections:

1. Title (use WRITE[title | Your Title Here])
2. Abstract (150-250 words summary)
3. Executive Summary (key findings and recommendations)
4. Introduction (context, importance, scope)
5. Background (relevant history and context)
6. Problem Statement (clear definition of the issue)
7. Solution/Methodology (your proposed approach)
8. Implementation (how to implement the solution)
9. Benefits (advantages and value proposition)
10. Challenges and Risks (potential issues and mitigation)
11. Recommendations (actionable next steps)
12. Conclusion (summary and future outlook)
13. References (citations and sources)
14. Appendices (additional information)

Tools available:
- PLAN[outline] - Create an outline/plan first
- SEARCH[query] - Search for information
- WRITE[section | content] - Write a section  
- SAVE_DOCX[filename] - Save the white paper to a DOCX file
- FINISH[done] - Complete the paper

WORKFLOW:
1. Start with ACTION: PLAN[list of sections]
2. Search and write sections
3. Call FINISH when done

IMPORTANT GUIDELINES:
- Research thoroughly using SEARCH before writing each major section
- Write detailed, well-researched content (minimum 200-300 words per section)
- Use professional, academic tone
- Include specific examples, data, and evidence
- Ensure logical flow between sections
- Only call FINISH after ALL sections are written

Format your response as:
THOUGHT: [Your reasoning]
ACTION: [One of the tools]

Begin by researching the topic!
"""
        current_prompt = system_prompt
        last_action = None
        repeat_count = 0
        
        for step in range(self.max_steps):
            print(f"\n--- Step {step + 1} ---")
            
            # Detect infinite loops - if same action repeats 3 times, stop
            if repeat_count >= 3:
                print("\n⚠️  Loop detected! Same action repeated 3 times.")
                print("Resetting and instructing agent to try different section...")
                repeat_count = 0  # Reset counter
                last_action = None
                # Add instruction to try something different
                current_prompt += "\nOBSERVATION: You're repeating the same action. Write a DIFFERENT section from the remaining list.\n"
                # Continue to next step instead of breaking
                continue
            
            # Rebuild prompt using context memory instead of unlimited growth
            if len(current_prompt) > 3000:
                print(f"🧠 Using context memory (prompt was {len(current_prompt)} chars)")
                context_summary = self.context_memory.get_context_summary()
                current_prompt = system_prompt + "\n\n" + context_summary + "\n\nContinue writing remaining sections.\n"
            
            response = self.generate_text(current_prompt, max_new_tokens=100)
            print(f"Agent Response:\n{response}")
            
            # More lenient parsing - look for any ACTION pattern
            action_patterns = [
                r'ACTION:\s*(SEARCH|WRITE|SAVE_DOCX|FINISH)\[([^\]]+)\]',
                r'(SEARCH|WRITE|SAVE_DOCX|FINISH)\[([^\]]+)\]',
            ]
            
            action_found = False
            for pattern in action_patterns:
                matches = re.finditer(pattern, response, re.IGNORECASE)
                for match in matches:
                    action_found = True
                    tool_name = match.group(1).upper()
                    tool_args = match.group(2).strip()
                    
                    print(f"🔧 Detected: {tool_name}[{tool_args[:50]}...]")
            
                    if tool_name == "PLAN":
                        try:
                            plan = tool_args.strip()
                            self.context_memory.set_plan(plan)
                            print(f"📋 Plan: {plan[:80]}...")
                            current_prompt += f"\nACTION: PLAN[{plan[:50]}...]\nOBSERVATION: Plan set. Now research and write.\n"
                        except Exception as e:
                            print(f"⚠️  Plan error: {e}")
                            current_prompt += f"\nOBSERVATION: Plan failed.\n"
                    
                    elif tool_name == "SEARCH":
                        try:
                            observation = self.search(tool_args)  # Already tracked in context memory
                            print(f"📚 Search result: {observation[:100]}...")
                            # Use summarized observation
                            current_prompt += f"\nACTION: SEARCH[{tool_args}]\nOBSERVATION: Found sources.\n"
                        except Exception as e:
                            print(f"⚠️  Search error: {e}")
                            current_prompt += f"\nOBSERVATION: Search failed: {e}\n"
                    
                    elif tool_name == "WRITE":
                        # Look for pipe separator
                        if "|" in tool_args:
                            parts = tool_args.split("|", 1)
                            section_title = parts[0].strip()
                            section_content = parts[1].strip() if len(parts) > 1 else ""
                            
                            # Remove any quotes from content
                            section_content = section_content.strip('"\'')
                            
                            # Detect if writing same section repeatedly
                            current_action = f"WRITE[{section_title}]"
                            if current_action == last_action:
                                repeat_count += 1
                            else:
                                repeat_count = 0
                            last_action = current_action
                            
                            # Check if already completed
                            if self.context_memory.should_skip_section(section_title):
                                print(f"⏭️  Skipping already completed: {section_title}")
                                current_prompt += f"\nOBSERVATION: '{section_title}' already written. Write a different section.\n"
                            else:
                                # Store section and track in memory
                                self.paper_content[section_title] = section_content
                                self.context_memory.add_section(section_title)
                                self.context_memory.add_action("WRITE", section_title)
                                print(f"✍️  Writing section: '{section_title}' ({len(section_content)} chars)")
                                current_prompt += f"\nACTION: WRITE[{section_title}]\nOBSERVATION: Section written.\n"
                        else:
                            print(f"⚠️  Missing | separator in WRITE[{tool_args[:30]}...]")
                            current_prompt += f"\nOBSERVATION: Use format: WRITE[section | content]\n"
                    
                    elif tool_name == "SAVE_DOCX":
                        try:
                            file_path = self.write_to_docx(tool_args)
                            print(f"💾 Saved to: {file_path}")
                            current_prompt += f"\nACTION: SAVE_DOCX[{tool_args}]\nOBSERVATION: Saved to {file_path}\n"
                        except Exception as e:
                            print(f"⚠️  Save error: {e}")
                            current_prompt += f"\nOBSERVATION: Save failed: {e}\n"
                    
                    elif tool_name == "FINISH":
                        print("🏁 Agent finishing...")
                        minimum_sections = 3  # Lowered for smaller model
                        if len(self.paper_content) < minimum_sections:
                            observation = f"Only {len(self.paper_content)} sections. Need at least {minimum_sections}."
                            print(f"⚠️  {observation}")
                            current_prompt += f"\nOBSERVATION: {observation}\n"
                        else:
                            # Auto-add references before finishing
                            if self.citations:
                                references = self.generate_references()
                                self.paper_content['References'] = references
                                print("\n📚 Auto-generated references section with {} citations".format(len(self.citations)))
                            
                            if self.paper_content:
                                self.write_to_docx()
                            print(f"\n✅ Complete! {len(self.paper_content)} sections written.")
                            return "White paper generation complete."
            
            if not action_found:
                print("⚠️  No valid action detected in response")
                current_prompt += "\nOBSERVATION: Use format: ACTION: PLAN, SEARCH, WRITE, or FINISH\n"

        return "Max steps reached."

@track_process_time
def generate(prompt):
    # This function is now a wrapper to start the agent
    agent = WhitePaperAgent(model, tokenizer)
    return agent.reason_and_act(prompt)

if __name__ == "__main__":
    model_path = "google/gemma-3-1b-it"
    model = load_pretrained_model(model_path)
    tokenizer = load_tokenizer_model(model_path)
    # model.eval() # Optional, depending on usage

    # You can now provide a detailed idea instead of just a topic
    idea = """
    Explore Adaptive Agent Selection for Multi-Agent LLM Systems, Focus on:

1) The agent will choose the best model based on the feedback given to each response when we routed the request to the different models the same query. Based on self agent as a judge rate and human feedback the models/agents which performing well should be used to repond to the queries.

2) This helps the best accurate results served to the user. Helps the trustworthy of the AI Usage in realtime usecases.

3) You can write a paper explaning how we can implement and achieve the solution

4) Highlight any potential risks and issues.

5) Write how it can help optimise the costs, and highlight if any other optimisation thoughts.
    """
    
    print(generate(idea))